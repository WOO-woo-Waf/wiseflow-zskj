# -*- coding: utf-8 -*-
import os
import random
import re
import time
import uuid
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from loguru import logger

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# 兼容不同版本的 python-docx：OxmlElement & qn
try:
    from docx.oxml.shared import OxmlElement, qn as qn_oxml
except ImportError:
    from docx.oxml import OxmlElement
    from docx.oxml.shared import qn as qn_oxml

# 第三方 / 业务依赖
from pb_api import PbTalker
from general_utils import get_logger_level

# OpenAI 客户端（兼容自定义 base_url 或仅 api_key）
from openai import OpenAI, RateLimitError


# ========== 环境 & 客户端 ==========
ROOT = Path(__file__).resolve().parents[2]
load_dotenv(ROOT / ".env", override=True)

base_url = os.environ.get("LLM_API_BASE", "")
token = os.environ.get("LLM_API_KEY", "")

if not base_url and not token:
    raise ValueError("LLM_API_BASE or LLM_API_KEY must be set")
elif base_url and not token:
    client = OpenAI(base_url=base_url)
elif not base_url and token:
    client = OpenAI(api_key=token)
else:
    client = OpenAI(api_key=token, base_url=base_url)

PROJECT_DIR = os.environ.get("PROJECT_DIR", "")
os.makedirs(PROJECT_DIR, exist_ok=True)

logger_file = os.path.join(PROJECT_DIR, "backend_service.log")
logger.add(
    logger_file,
    level=get_logger_level(),
    backtrace=True,
    diagnose=True,
    rotation="50 MB",
)

pb = PbTalker(logger)

# LLM & 输入大小提示（可根据所用模型调整）
REPORT_MODEL = os.environ.get("REPORT_MODEL", "gpt-4o-mini-2024-07-18")
MAX_ITEM_CHARS = 10000          # 单条原始材料截断
MAX_ABSTRACT_CHARS = 4000       # 单篇文章摘要截断
MAX_ARTICLES_PER_ITEM = 30      # 每条编号下材料里最多塞几篇文章摘要


# ========== 固定模板常量 ==========
SECTIONS = [
    ("综合要闻", "general"),
    ("区域新闻", "regional"),
    ("政策数据", "policy"),
    ("科技前沿", "tech"),
    ("行业动态", "industry"),  # 含子类
    ("对标资讯", "benchmark"),
    ("中核要闻", "cnnc_headline"),
]
INDUSTRY_SUB = ["核能", "清洁能源", "环保", "金融", "核技术应用", "智能信息", "其他"]

TAG_MAP = {
    # 一级
    "综合要闻": ("general", None),
    "区域新闻": ("regional", None),
    "政策数据": ("policy", None),
    "科技前沿": ("tech", None),
    "对标资讯": ("benchmark", None),
    "中核要闻": ("cnnc_headline", None),
    # 行业动态子类
    "核能": ("industry", "核能"),
    "清洁能源": ("industry", "清洁能源"),
    "环保": ("industry", "环保"),
    "金融": ("industry", "金融"),
    "核技术应用": ("industry", "核技术应用"),
    "智能信息": ("industry", "智能信息"),
}


# ========== 工具函数 ==========
def cn_today_str(dt: datetime | None = None) -> str:
    dt = dt or datetime.now()
    return f"{dt.year}年{dt.month}月{dt.day}日"


def safe_filename(name: str) -> str:
    name = (name or "").strip()
    name = re.sub(r'[\/\\\:\*\?"<>\|]', "_", name)
    name = re.sub(r"\s+", " ", name)
    return name or f"中核日报（{cn_today_str()}）"


def openai_llm(messages: list, model: str, logger_=None, **kwargs) -> str:
    """
    带指数退避的健壮 LLM 调用：
    - 对 504 / 超时 / 临时网络错误 等进行最多 5 次重试
    - 默认每次超时 60s（可通过 kwargs['timeout'] 覆盖）
    - 抖动重试间隔：2s, 4s, 8s, 12s, 16s（±0.5s）
    """
    import random, time

    max_retries = int(kwargs.pop("max_retries", 5))
    base_delay = 2.0
    timeout = kwargs.pop("timeout", 60)

    if logger_:
        logger_.debug(f"messages:\n {messages}")
        logger_.debug(f"model: {model}")
        logger_.debug(f"kwargs:\n {kwargs}")

    def _should_retry(exc: Exception) -> bool:
        s = (str(exc) or "").lower()
        keys = ["504", "gateway", "time-out", "timeout", "temporarily", "connection", "reset", "unavailable"]
        return any(k in s for k in keys)

    last_err = None
    for attempt in range(max_retries):
        try:
            resp = client.chat.completions.create(
                messages=messages,
                model=model,
                timeout=timeout,
                **kwargs
            )
            if not getattr(resp, "choices", None):
                if logger_: logger_.warning(f"openai_llm warning: empty choices: {resp}")
                last_err = RuntimeError("empty choices")
                raise last_err
            if logger_:
                logger_.debug(f"result:\n {resp.choices[0]}")
                if getattr(resp, "usage", None):
                    logger_.debug(f"usage:\n {resp.usage}")
            return resp.choices[0].message.content
        except Exception as e:
            last_err = e
            if not _should_retry(e) or attempt == max_retries - 1:
                if logger_: logger_.error(f"openai_llm error (no more retries): {e}")
                break
            delay = base_delay * (2 ** attempt) + random.uniform(-0.5, 0.5)
            if logger_: logger_.warning(f"{e}\nretrying in {delay:.1f}s (attempt {attempt+1}/{max_retries})")
            time.sleep(max(0.5, delay))

    return ""



def add_hyperlink(paragraph, url, text):
    """在段落中插入一个可点击超链接（蓝色+下划线）"""
    part = paragraph.part
    try:
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    except Exception:
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn_oxml("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn_oxml("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn_oxml("w:val"), "0000FF")

    rPr.append(u)
    rPr.append(color)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _norm_date(d):
    d = str(d or "")
    digits = re.sub(r"[^\d]", "", d)
    if len(digits) >= 8:
        return f"{digits[:4]}-{digits[4:6]}-{digits[6:8]}"
    return ""


def _collect_keywords(insights: list[dict]) -> list[str]:
    seen, res = set(), []
    for it in insights or []:
        kws = it.get("keywords") or []
        if isinstance(kws, str):
            kws = [x.strip() for x in re.split(r"[，、,\s]+", kws) if x.strip()]
        for k in kws:
            if k not in seen:
                seen.add(k)
                res.append(k)
    return res


def _load_role_config():
    """从 PB 获取角色设定（roleplays 集合）"""
    try:
        role_cfg = pb.read(collection_name="roleplays", filter="activated=True")
        if role_cfg:
            character = role_cfg[0].get("character", "") or ""
            report_type = role_cfg[0].get("report_type", "") or ""
            return character, report_type, role_cfg[0].get("id", "")
    except Exception as e:
        logger.warning(f"load roleplays failed: {e}")
    return "", "", ""

CN_TO_KEY = {cn: key for cn, key in SECTIONS}
SECTION_KEYS = {key for _, key in SECTIONS}

# 分隔符：/ - ｜ | ： : 以及中英文空白与常见标点
_SPLIT_RE = re.compile(r"[\s　/／\-–—\|｜:：,，、()（）\.。]+")
# 仅用于“等值”比较时的去标点归一
def _norm_eq(s: str) -> str:
    return re.sub(r"[\s　:：\-/|（）()、,，.。]+", "", s or "")

def classify_item(tag: str | None, category: str | None = None):
    """
    只看 category：
      - 若为标准大类（等值匹配/去标点等值），直接归入该大类：(key, None)
      - 若包含“行业动态”，按分隔符拆分，取其后的第一段作为子类：(industry, 子类或None)
      - 其他情况严格兜底到“综合要闻”：("general", None)
    """
    cat = (category or "").strip()
    # 1) 直接中文大类命中
    if cat in CN_TO_KEY:
        key = CN_TO_KEY[cat]
        return (key, None) if key != "industry" else ("industry", None)

    # 2) 直接英文 key 命中（如果有人直接传 key）
    if cat in SECTION_KEYS:
        return (cat, None) if cat != "industry" else ("industry", None)

    # 3) 去标点后的等值匹配（仍是“等值”，不做猜测）
    ncat = _norm_eq(cat)
    for cn, key in SECTIONS:
        if ncat and ncat == _norm_eq(cn):
            return (key, None) if key != "industry" else ("industry", None)

    # 4) 行业动态：拆分并取子类
    #    支持：行业动态/核能、行业动态-核能、行业动态｜核能、industry/xxx
    if cat:
        parts = [p for p in _SPLIT_RE.split(cat) if p]  # 去空、去空段
        # 找到“行业动态”或 industry 的位置
        idx = -1
        for i, p in enumerate(parts):
            if p == "行业动态" or p.lower() == "industry":
                idx = i
                break
        if idx != -1:
            sub = None
            if idx + 1 < len(parts):
                sub = parts[idx + 1] or None
            return ("industry", sub)

    # 5) 兜底严格回落“综合要闻”
    return ("general", None)


# ========== 逐洞见建模（三件套） ==========
def _extract_article_summaries(ent: dict) -> list[dict]:
    """从 ent['articles'] 抽取标题/日期/摘要/URL；长度兜底"""
    arts = []
    for a in (ent.get("articles") or [])[:MAX_ARTICLES_PER_ITEM]:
        arts.append({
            "title": (a.get("title") or "").strip(),
            "date": _norm_date(a.get("publish_time")),
            "abstract": (a.get("abstract") or a.get("content") or "").strip()[:MAX_ABSTRACT_CHARS],
            "url": (a.get("url") or "").strip(),
        })
    return arts


def _recent_time(arts: list[dict]) -> str:
    ts = [x.get("date") for x in arts if x.get("date")]
    return max(ts) if ts else ""


def _process_insight_item(ent: dict, character: str, report_type: str) -> dict:
    """
    输入：{'content': 洞见摘要, 'articles': [...], 'url': 源链接, 可选 'title','category','tag','keywords'}
    输出：{'title','summary','sources'[1~3 urls],'time'}
    """
    content = (ent.get("content") or "").strip()[:MAX_ITEM_CHARS]
    articles = _extract_article_summaries(ent)
    recent_time = _recent_time(articles)

    # 源链接优先：洞见本身的 url
    src_urls = []
    if isinstance(ent.get("url"), str) and ent["url"].strip():
        src_urls.append(ent["url"].strip())
    # 再补充文章里的 URL
    for a in articles:
        if a.get("url") and a["url"] not in src_urls:
            src_urls.append(a["url"])
    src_urls = src_urls[:3]

    # 提示词（参考你给的写法，改为 OpenAI）
    character = character or "专业情报分析官"
    report_type = report_type or "综合情报"
    sys = (
        f"你是一名{character}。根据“洞见摘要”和“原始文章摘录”，输出结构化结果：\n"
        "1) concise_title：50字内提要式标题；\n"
        "2) detailed_summary：120~500字，准确客观、书面化；\n"
        "3) sources：最多3个 URL（原样返回）。\n"
        "必须严格输出 JSON。"
    )
    usr = {
        "report_type": report_type,
        "insight_summary": content,
        "primary_url": ent.get("url", ""),
        "articles": articles
    }

    title = ""; summary = ""; sources = []
    try:
        out = openai_llm(
            messages=[{"role": "system", "content": sys},
                      {"role": "user", "content": str(usr)}],
            model=REPORT_MODEL,
            temperature=0.2,
            logger_=logger
        ) or ""
        m = re.search(r"\{.*\}", out, flags=re.S)
        if m:
            import json
            data = json.loads(m.group(0))
            title = (data.get("concise_title") or "").strip()
            summary = (data.get("detailed_summary") or "").strip()
            s = data.get("sources") or []
            if isinstance(s, str):
                s = [s]
            sources = [x.strip() for x in s if isinstance(x, str) and x.strip()]
    except Exception as e:
        logger.warning(f"_process_insight_item llm fail: {e}")

    # 兜底
    if not title:
        title = (ent.get("title") or (articles[0]["title"] if articles and articles[0]["title"] else content[:24])).strip()
    if not summary:
        base = content
        if articles and articles[0].get("abstract"):
            base = f"{base}\n{articles[0]['abstract']}"
        summary = base[:180]
    if not sources:
        sources = src_urls

    return {"title": title, "summary": summary, "sources": sources[:3], "time": recent_time}

# ========== 分类聚合 & 类别内逻辑排序 ==========
def _group_by_section(insights: list[dict], character: str, report_type: str) -> dict:
    """
    返回：{section_key: {'title':中文名, 'items':[processed...], 'subs':{子类:[processed...]}}}
    processed: {'title','summary','sources','time','_raw'}
    """
    grouped = {
        key: {'title': title, 'items': [], 'subs': {s: [] for s in INDUSTRY_SUB}}
        for title, key in SECTIONS
    }

    for ent in insights or []:
        sec_key, sub = classify_item(ent.get("tag"), ent.get("category"))
        proc = _process_insight_item(ent, character, report_type)
        proc["_raw"] = ent
        if sec_key == "industry":
            grouped[sec_key]['subs'][sub or "其他"].append(proc)
        else:
            grouped[sec_key]['items'].append(proc)
    return grouped

# ========== 先粗分“原始洞见”到各分区/子类（不做逐条LLM） ==========
def _group_raw_by_category(insights: list[dict]) -> dict:
    """
    返回：{section_key: {'title':中文名, 'raw':[...], 'subs':{子类:[...]}}}
    注意：这里只按 ent['category']/‘目录’/‘section’/‘sec’ 放桶，不做 LLM 处理。
    """
    grouped = {
        key: {'title': title, 'raw': [], 'subs': {s: [] for s in INDUSTRY_SUB}}
        for title, key in SECTIONS
    }

    # 计数与示例收集（调试用）
    per_bucket_count = {key: 0 for _, key in SECTIONS}
    samples = []

    for idx, ent in enumerate(insights or []):
        # category 支持多种字段名
        category = (ent.get("category")
                    or ent.get("目录")
                    or ent.get("section")
                    or ent.get("sec")
                    or "")
        sec_key, sub = classify_item(None, category)

        # —— 仅把“行业动态/子类”用到 subs；其他分区全部进 raw
        if sec_key == "industry":
            grouped[sec_key]['subs'][sub or "其他"].append(ent)
        else:
            grouped[sec_key]['raw'].append(ent)

        per_bucket_count[sec_key] = per_bucket_count.get(sec_key, 0) + 1

        # 保存样例（最多前 10 条），便于追踪错分
        if len(samples) < 10:
            samples.append({
                "i": idx,
                "id": ent.get("id"),
                "category_raw": category,
                "classified": f"{sec_key}" + (f"/{sub}" if sec_key == "industry" else ""),
            })

    # —— 输出调试汇总
    try:
        logger.info("[_group_raw_by_category] bucket counts: " + ", ".join(
            f"{k}={per_bucket_count.get(k,0)}" for _, k in SECTIONS
        ))
        if samples:
            logger.debug("[_group_raw_by_category] first-samples: " + "; ".join(
                f"#{s['i']} id={s['id']} cat={s['category_raw']!r} -> {s['classified']}"
                for s in samples
            ))
    except Exception:
        pass

    return grouped


def process_category_batch(
    category_cn: str,
    entries: list[dict],
    character: str,
    report_type: str,
    max_articles_per_item: int = 5,
) -> list[dict]:
    """
    批量处理一个类别下的所有洞见；按批（每批<=20）与 LLM 交互；
    合并所有批次结果后再做全局排序（YYYY-MM-DD 倒序；缺失在后，稳定）。
    """
    # ===== 内部工具 =====
    BATCH_LIMIT = 10  # 每批最多 20 条

    def _mk_payload_item(ent: dict) -> dict:
        content = (ent.get("content") or "").strip()[:MAX_ITEM_CHARS]
        arts_all = _extract_article_summaries(ent)  # [{'title','date','abstract','url'}]
        arts = arts_all[:max_articles_per_item]
        recent_time = _recent_time(arts_all)

        urls = []
        if isinstance(ent.get("url"), str) and ent["url"].strip():
            urls.append(ent["url"].strip())
        for a in arts:
            u = (a.get("url") or "").strip()
            if u and u not in urls:
                urls.append(u)
        urls = urls[:3]

        for a in arts:
            if a.get("abstract"):
                a["abstract"] = a["abstract"][: min(400, MAX_ABSTRACT_CHARS)]

        return {
            "id": ent.get("id") or "",
            "content": content,
            "primary_url": (ent.get("url") or ""),
            "recent_time": recent_time,
            "articles": arts,
            "fallback": {
                "title": (ent.get("title") or (arts_all[0]["title"] if arts_all and arts_all[0].get("title") else content[:24])).strip(),
                "summary_src": (content + ("\n" + arts_all[0]["abstract"] if (arts_all and arts_all[0].get("abstract")) else "")),
                "urls": urls,
            },
        }

    def _fallback_sorted_items(prepared: list[dict]) -> list[dict]:
        res = []
        for p in prepared:
            fb = p["fallback"]
            res.append({
                "title": fb["title"],
                "summary": (fb["summary_src"][:180]).strip(),
                "sources": fb["urls"],
                "time": p.get("recent_time") or "",
            })
        # 按 time 倒序；无时间排后；稳定排序
        def _key(x):
            t = (x.get("time") or "").strip()
            return (0, t) if t else (1, "")
        res.sort(key=_key, reverse=True)
        return res

    def _split_batches(seq: list, n: int) -> list[list]:
        return [seq[i:i+n] for i in range(0, len(seq), n)]

    def _post_sort_global(items: list[dict]) -> list[dict]:
        # 全局排序：YYYY-MM-DD 倒序，无时间的在后；保持稳定
        def _key(x):
            t = (x.get("time") or "").strip()
            return (0, t) if t else (1, "")
        return sorted(items, key=_key, reverse=True)

    # ===== 空输入直接返回 =====
    if not entries:
        return []

    character = character or "专业情报分析官"
    report_type = report_type or "综合情报"

    # 预处理：紧凑化所有条目（一次）
    prepared_all = [_mk_payload_item(e) for e in entries]
    id2prepared = {p["id"]: p for p in prepared_all}

    # 按批切分
    batches = _split_batches(prepared_all, BATCH_LIMIT)

    all_results: list[dict] = []

    # 逐批调用 LLM（若该批失败，则对该批走兜底）
    for batch_idx, prepared in enumerate(batches, 1):
        sys_prompt = (
            f"你是一名{character}。现在给你同一类别（{category_cn}）下的多条洞见的精简材料，"
            "请你为每条洞见生成：\n"
            "  - concise_title：<=50字的提要式标题；\n"
            "  - detailed_summary：2000~500字，准确客观、书面化；\n"
            "  - sources：最多3个URL（仅在材料出现过的 URL 中选择，不得臆造）；\n"
            "  - time：若能从材料判断出最近的时间，按 YYYY-MM-DD 返回；无法判断留空。\n"
            "然后对这些条目进行【逻辑排序】（同一主题聚合、政策/权威优先、时间倒序等可综合判断）。\n"
            "严格输出 JSON，格式：\n"
            '{ "items":[ {"id":"原样返回","concise_title":"...","detailed_summary":"...","sources":["..."],"time":"YYYY-MM-DD"}, ... ] }\n'
            "注意：\n"
            "1) 禁止杜撰事实或链接；\n"
            "2) sources 仅可来自材料里的 primary_url 或 articles.url；\n"
            "3) items 中的 id 必须与输入一一对应；\n"
            "4) 允许合并主题相近的条目，但不要丢失关键信息；\n"
            "5) 若无法合并，请保持每条都输出。\n"
        )
        usr_payload = {
            "report_type": report_type,
            "category": category_cn,
            "records": [
                {
                    "id": p["id"],
                    "insight_summary": p["content"],
                    "primary_url": p["primary_url"],
                    "recent_time": p["recent_time"],
                    "articles": p["articles"],
                }
                for p in prepared
            ]
        }

        try:
            out = openai_llm(
                messages=[{"role": "system", "content": sys_prompt},
                          {"role": "user", "content": str(usr_payload)}],
                model=REPORT_MODEL,
                temperature=0.2,
                logger_=logger,
                timeout=60,
                max_retries=5,
            ) or ""

            import json
            m = re.search(r"\{.*\}", out, flags=re.S)
            data = json.loads(m.group(0)) if m else None
            if not data or "items" not in data or not isinstance(data["items"], list):
                raise ValueError("invalid LLM JSON")

            prepared_by_id = {p["id"]: p for p in prepared}
            batch_results: list[dict] = []

            for it in data["items"]:
                _id = (it.get("id") or "").strip()
                base = prepared_by_id.get(_id)
                if not base:
                    continue

                # sources：仅保留材料中允许的 URL
                allowed = set(base["fallback"]["urls"])
                srcs = it.get("sources") or []
                if isinstance(srcs, str):
                    srcs = [srcs]
                filtered_sources = []
                for s in srcs:
                    s = (s or "").strip()
                    if s and s in allowed and s not in filtered_sources:
                        filtered_sources.append(s)
                    if len(filtered_sources) >= 3:
                        break
                if not filtered_sources:
                    filtered_sources = base["fallback"]["urls"]

                title = (it.get("concise_title") or "").strip()
                summary = (it.get("detailed_summary") or "").strip()
                time_str = (it.get("time") or "").strip()

                if not title:
                    title = base["fallback"]["title"]
                if not summary:
                    ss = base["fallback"]["summary_src"]
                    summary = ss[: max(180, min(500, len(ss)))]

                batch_results.append({
                    "title": title,
                    "summary": summary,
                    "sources": filtered_sources[:3],
                    "time": time_str or base.get("recent_time") or "",
                })

            # 批次产出过少 → 该批走兜底
            if len(batch_results) < max(1, len(prepared) // 2):
                raise ValueError("too few items from LLM in this batch, use fallback")

            all_results.extend(batch_results)

        except Exception as e:
            if 'logger' in globals() and logger:
                logger.warning(f"[process_category_batch] batch#{batch_idx} LLM fail, fallback used: {e}")
            all_results.extend(_fallback_sorted_items(prepared))

    # 全部批次合并后做一次全局排序
    return _post_sort_global(all_results)


def _logical_sort_items_via_llm(section_title: str, items: list[dict]) -> list[dict]:
    """用 LLM 给出类别内排序；失败回退时间倒序"""
    if not items:
        return items

    payload = [{"idx": i + 1, "title": it.get("title", ""), "summary": it.get("summary", ""), "time": it.get("time", "")}
               for i, it in enumerate(items)]

    sys = (
        "你是资深编辑。请对同一类别的新闻条目做逻辑排序："
        "优先级=政策/监管>项目工程>企业行动>数据/研究；同主题聚在一起，时间相近相邻。"
        "只输出 JSON：{\"order\":[原始idx,...]}。"
    )
    order = []
    try:
        out = openai_llm(
            messages=[{"role": "system", "content": sys},
                      {"role": "user", "content": str({"section": section_title, "items": payload})}],
            model=REPORT_MODEL,
            temperature=0.1,
            logger_=logger
        ) or ""
        m = re.search(r"\{.*\}", out, flags=re.S)
        if m:
            import json
            data = json.loads(m.group(0))
            order = [int(x) for x in (data.get("order") or []) if isinstance(x, int)]
    except Exception as e:
        logger.warning(f"logical sort llm error: {e}")

    if order and set(order) == set(range(1, len(items) + 1)):
        return [items[i - 1] for i in order]

    # 兜底：时间倒序
    def _k(it): return it.get("time") or "", it.get("title") or ""
    return sorted(items, key=_k, reverse=True)


PB_BASE_URL = os.environ.get("PB_BASE_URL", "").rstrip("/")

def _pb_file_url(collection: str, record_id: str, filename: str) -> str:
    """
    生成 PocketBase 文件下载链接（无需 token 的公共文件可直接访问；
    若你配置了鉴权下载，这里也可以改为调用 pb 提供的 file_url 方法）。
    """
    if not (PB_BASE_URL and collection and record_id and filename):
        return ""
    return f"{PB_BASE_URL}/api/files/{collection}/{record_id}/{filename}"

def _save_report_memory(title: str, snapshot_text: str, docx_local_path: str) -> tuple[str, str]:
    """
    在 PB.report_memories 创建一条记录 + 上传 DOCX 到字段 docx，并将 docx_path 写回。
    返回: (memory_id, docx_path)
    """
    try:
        # 1) 先创建记录（不含文件）
        mem_id = pb.add(
            collection_name="report_memories",
            body={
                "title": title or "",
                "snapshot": snapshot_text or "",
                "docx_path": "",      # 占位，上传后回写
            },
        )
        mem_id = str(mem_id)

        # 2) 上传 DOCX 文件到这条记录的 docx 字段
        filename = os.path.basename(docx_local_path) if docx_local_path else f"{title or '报告'}.docx"
        with open(docx_local_path, "rb") as f:
            _ = pb.upload("report_memories", mem_id, "docx", filename, f)

        # 3) 生成下载地址（或用 pb.file_url(...) 如你封装了）
        docx_path = _pb_file_url("report_memories", mem_id, filename)

        # 4) 回写 docx_path
        _ = pb.update("report_memories", mem_id, {"docx_path": docx_path})

        return mem_id, docx_path
    except Exception as e:
        logger.error(f"_save_report_memory failed: {e}")
        return "", ""


def _update_report_memory(jawbone_id: str, title: str, snapshot_text: str, docx_path: str) -> None:
    body = {
        "title": title,
        "snapshot": snapshot_text,
        "docx_path": docx_path,
        "updated": datetime.now().isoformat(timespec="seconds")
    }
    try:
        pb.update(collection_name="report_memories", id=jawbone_id, body=body)
    except Exception as e:
        logger.warning(f"update memory failed: {e}")

def _get_report_memory_by_id(memory_id: str) -> dict | None:
    try:
        recs = pb.read("report_memories",
                       fields=["id", "title", "snapshot", "docx_path", "created", "updated"],
                       filter=f'id="{memory_id}"')
        return recs[0] if recs else None
    except Exception as e:
        logger.warning(f"read report memory failed: {e}")
        return None
    

def fetch_active_tags_from_pb() -> list[str]:
    # 只要激活的标签，按更新时间倒序
    recs = pb.read(
        "tags",filter="activated=true" 
    ) or []
    return [r.get("name", "").strip() for r in recs if r.get("name")]

# ========== LLM 报告生成（首次 / 或按记忆改写） ==========
def get_report(
    insight_entries: list[dict],
    articles: list[dict],
    memory: str,         # 若提供，则走“改写”流程
    topics: list[str],
    comment: str,
    docx_file: str
) -> tuple[bool, str, str]:
    """
    两种模式：
    1) 首次生成：基于多洞见→逐条建模→分类→排序→组装→DOCX→入库
    2) 改写：有 memory + comment 时，只按意见改写 memory→DOCX→更新入库
    返回: (ok, snapshot_text, report_title)
    """
    character, report_type, _role_id = _load_role_config()
    today = cn_today_str()
    report_title = topics[0].strip() if (topics and isinstance(topics, list) and (topics[0] or "").strip()) else f"中核日报（{today}）"

    # ========== 改写模式 ==========
    if (memory or "").strip() and (comment or "").strip():
        logger.debug("rewrite mode with memory + comment")
        new_text = revise_snapshot_text(memory, comment, logger_=logger)
        if not new_text:
            return False, "", report_title
        # DOCX
        ok = build_docx_from_snapshot(
            snapshot_text=new_text,
            articles=articles,
            docx_file=docx_file,
            always_appendix=True,
            inline_links=True,
            grouped_for_links=None
        )
        # 保存/更新记忆
        id, _docx_path = _save_report_memory(report_title, new_text, docx_file)
        return ok, new_text, report_title, id

    # ========== 首次生成 ==========
    # # 1) 分类 + 逐洞见建模
    # grouped = _group_by_section(insight_entries, character, report_type)

    # # 2) 类别内逻辑排序
    # for title_cn, key in SECTIONS:
    #     if key == "industry":
    #         for sub in INDUSTRY_SUB:
    #             grouped[key]['subs'][sub] = _logical_sort_items_via_llm(f"{title_cn}-({sub})", grouped[key]['subs'][sub])
    #     else:
    #         grouped[key]['items'] = _logical_sort_items_via_llm(title_cn, grouped[key]['items'])

     # 1) 仅按目录粗分原始洞见
    grouped_raw = _group_raw_by_category(insight_entries)

    # 2) 类别内“一次性批处理”（模型排序；失败兜底按时间倒序）
    grouped_processed = {
        key: {'title': grouped_raw[key]['title'], 'items': [], 'subs': {s: [] for s in INDUSTRY_SUB}}
        for _, key in SECTIONS
    }
    for title_cn, key in SECTIONS:
        if key != "industry":
            raw_items = grouped_raw[key]['raw']
            grouped_processed[key]['items'] = process_category_batch(title_cn, raw_items, character, report_type)
        else:
            for sub in INDUSTRY_SUB:
                raw_items = grouped_raw[key]['subs'][sub]
                grouped_processed[key]['subs'][sub] = process_category_batch(f"{title_cn}（{sub}）", raw_items, character, report_type)

    # 3) 关键词：从 PB.tags 读取激活标签
    try:
        kws = fetch_active_tags_from_pb()
    except Exception as e:
        logger.warning(f"fetch_active_tags_from_pb failed: {e}")
        kws = []

    # 4) 组装最终正文（分区标题不加序号；每条=编号行→概括→链接行）
    lines = [report_title]
    if kws:
        lines.append(f"关键词：{'、'.join(kws)}")

    def _emit_item_block(idx: int, it: dict) -> list[str]:
        urls = [u for u in (it.get("sources") or []) if u][:3]
        blk = [f"{idx}，{(it.get('title') or '').strip()}"]
        if it.get("summary"):
            blk.append(it["summary"].strip())
        blk.extend(urls)
        return blk

    has_any = False
    for title_cn, key in SECTIONS:
        if key != "industry":
            items = grouped_processed[key]['items']          # ★ 用 processed
            if not items:
                continue
            has_any = True
            lines.append(f"{title_cn}：")
            for i, it in enumerate(items, start=1):
                lines.extend(_emit_item_block(i, it))
        else:
            subs_have = any(grouped_processed[key]['subs'][s] for s in INDUSTRY_SUB)
            if not subs_have:
                continue
            has_any = True
            lines.append(f"{title_cn}：")
            for sub in INDUSTRY_SUB:
                items = grouped_processed[key]['subs'][sub]  # ★ 用 processed
                if not items:
                    continue
                lines.append(f"（{sub}）")
                for i, it in enumerate(items, start=1):
                    lines.extend(_emit_item_block(i, it))

    snapshot_text = "\n".join(lines)
    if not has_any:
        logger.warning("no content after grouping")
        return False, "", report_title

    # 5) DOCX 渲染
    ok = build_docx_from_snapshot(
        snapshot_text=snapshot_text,
        articles=articles or [],
        docx_file=docx_file,
        always_appendix=True,
        inline_links=True,
        grouped_for_links=None  # URL 已在正文中
    )
    if not ok:
        return False, "", report_title

    # 6) 入库记忆
    id, _docx_path = _save_report_memory(report_title, snapshot_text, docx_file)

    return ok, snapshot_text, report_title, id


# ========== LLM 结构锁定改写（分块版） ==========

import re
from typing import List, Tuple

def _detect_header(lines: List[str]) -> Tuple[str | None, str | None, int]:
    """
    从文本行里识别：标题行 + 关键词行（若存在）。
    返回: (title_line, keywords_line, body_start_index)
    """
    i = 0
    # 跳过起始空行
    while i < len(lines) and not lines[i].strip():
        i += 1
    title = lines[i] if i < len(lines) else None
    i += 1
    keywords = None
    if i < len(lines) and lines[i].strip().startswith("关键词："):
        keywords = lines[i]
        i += 1
    return title, keywords, i

_SECTION_RE = re.compile(r"^.+?：$")   # 分区标题（末尾中文冒号）
_NUMBER_RE  = re.compile(r"^\d+，")    # 编号行（形如“1，”开头）

def _is_boundary_line(s: str) -> bool:
    s = s.strip()
    return bool(_SECTION_RE.match(s) or _NUMBER_RE.match(s) or not s)

def _chunk_lines(lines: List[str], max_chars: int) -> List[List[str]]:
    """
    将正文行按“逼近 max_chars”进行分块；优先在“分区标题/编号/空行”边界断开。
    """
    chunks: List[List[str]] = []
    cur: List[str] = []
    cur_len = 0
    last_boundary_idx = -1  # 记录当前块内最近的边界位置（相对 cur 的下标）

    for ln in lines:
        add_len = len(ln) + 1  # +换行
        # 若加上本行超过上限，则在最近边界处切块；若没有边界则硬切
        if cur and cur_len + add_len > max_chars:
            cut_at = last_boundary_idx if last_boundary_idx >= 0 else len(cur) - 1
            chunks.append(cur[:cut_at + 1])
            rest = cur[cut_at + 1:]
            cur = rest[:]  # 剩余行续入下一块
            cur_len = sum(len(x) + 1 for x in cur)
            last_boundary_idx = -1
            # 重置后再尝试加入当前行
        # 累加当前行
        cur.append(ln)
        cur_len += add_len
        if _is_boundary_line(ln):
            last_boundary_idx = len(cur) - 1

    if cur:
        chunks.append(cur)
    return chunks

# 允许两种冒号（：/:），允许标题行末尾有空格；允许编号使用 ，、,.． 及前后空格
_SECTION_RE = re.compile(r"^\s*([^\n：:]+?)\s*[：:]\s*$", re.M)
_NUMBER_RE  = re.compile(r"^\s*\d+\s*[，,、\.．]\s*", re.M)

def _pre_norm(txt: str) -> str:
    """仅用于结构计数的宽松归一：去掉行尾空白，统一换行。"""
    if not txt:
        return ""
    txt = txt.replace("\r\n", "\n").replace("\r", "\n")
    # 去除 Markdown 强制换行产生的行尾空格
    txt = re.sub(r"[ \t]+\n", "\n", txt)
    return txt

def _struct_count(txt: str) -> tuple[int, int]:
    t = _pre_norm(txt)
    secs = _SECTION_RE.findall(t)
    nums = _NUMBER_RE.findall(t)
    return (len(secs), len(nums))

def _call_llm_revise(piece_text: str, comment: str, idx: int, total: int, has_header: bool, logger_):
    """
    对单个片段调用 LLM 改写；对首块（含标题与关键词）要求保留，对非首块禁止重复标题/关键词。
    """
    sys_prompt = (
        "你是报告改写助手。请在【严格保留结构与编号】的前提下，仅改写“正文行”，其它行必须逐字保留：\n"
        "【文档结构与术语】\n"
        "• 分区标题行：形如“综合要闻：”/“政策数据：”，末尾是冒号。不可修改。\n"
        "• 子类标头行：形如“（核能）”。不可修改。\n"
        "• 关键词行：形如“关键词：……”。不可修改。\n"
        "• 条目单元：由三部分组成（顺序固定）：\n"
        "  A. 标题行：以“数字+全角逗号”开头（如“1，”），通常包含条目标题。可修改。\n"
        "  B. 正文行：紧随标题行，直到链接行之前的 1~N 行文本。可修改】。\n"
        "  C. 链接行：以 http/https 开头的 1~N 行。不可修改、不可增删、不可改序。\n"
        "【必须遵守】\n"
        "1) 仅改写“正文行”；其余所有行（分区标题/子类标头/关键词/标题行/链接行）必须与输入完全一致（字符级一致）。\n"
        "2) 保证行序不变。\n"
        "3) 不得新增或删除链接，不得改写任何链接字符（含大小写、协议、查询串、末尾斜杠等）。\n"
        "4) 保持编号与标点完全一致（如“1，”仍为“1，”）；禁止把“，”改成“、/./,”等；禁止在行尾添加两个空格或其它 Markdown 格式。\n"
        "5) 保持事实准确，不得臆造；若修改意见与事实冲突，优先维持原文。\n"
        "6) 若任何规则与输出产生冲突，请原样输出输入文本（即不做任何修改）。\n"
        "【输出格式】\n"
        "• 仅输出改写后的完整片段纯文本；不要附加解释、Markdown 标记或引号。\n"
    )

    guide = (
        f"【片段 {idx}/{total}；{'含' if has_header else '不含'}标题/关键词】\n"
        "请严格按上述规则处理：只改写“正文行”，其他所有行逐字照抄。"
    )

    usr = (
        guide
        + "\n\n【修改意见】\n"
        + comment
        + "\n\n【片段原文】\n"
        + piece_text
        + "\n\n【请输出】\n"
        "输出与输入行数一致；仅正文行发生文字差异，其它行必须完全相同。"
    )

    out = openai_llm(
        messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": usr}],
        model=REPORT_MODEL,
        temperature=0.2,
        logger_=logger_
    ) or ""
    return out

def revise_snapshot_text_chunked(
    snapshot_text: str,
    comment: str,
    logger_=None,
    max_chunk_chars: int = 6000,  # 可按模型上下文调整
) -> str:
    """
    按行分块，逐块改写并合并。首块包含标题与“关键词：”，其余块只含正文。
    """
    if not snapshot_text or not (comment or "").strip():
        return ""

    lines = snapshot_text.splitlines()
    title, keywords, body_start = _detect_header(lines)
    body_lines = lines[body_start:]

    # 原结构统计（用于最终校验）
    orig_secs, orig_nums = _struct_count(snapshot_text)

    # 分块：正文行按 max_chunk_chars 切；首块在输出时会补上标题/关键词（若存在）
    pieces = _chunk_lines(body_lines, max_chunk_chars)
    total = len(pieces) if pieces else 1

    revised_chunks: List[str] = []

    for idx, piece in enumerate(pieces or [[]], start=1):
        # 组装本片段输入文本
        has_header = (idx == 1 and (title or keywords))
        part_lines = []
        if has_header:
            if title:    part_lines.append(title)
            if keywords: part_lines.append(keywords)
        part_lines.extend(piece)
        piece_text = "\n".join(part_lines).rstrip()

        try:
            out = _call_llm_revise(piece_text, comment, idx, total, has_header, logger_)
        except Exception as e:
            if logger_:
                logger_.error(f"revise chunk#{idx} failed: {e}")
            out = ""

        # 结构对齐校验：片段内“分区标题/编号行”数量要与输入片段一致
        in_secs, in_nums   = _struct_count(piece_text)
        out_secs, out_nums = _struct_count(out)
        if (in_secs, in_nums) != (out_secs, out_nums) or not out.strip():
            # 回退：使用原始片段
            if logger_:
                logger_.warning(f"revise chunk#{idx} rejected due to structure mismatch or empty output")
            out = piece_text

        revised_chunks.append(out.strip())

    merged = "\n".join(revised_chunks).strip()

    # 最终结构校验与回退（与原文比较）
    if _struct_count(merged) != (orig_secs, orig_nums):
        if logger_:
            logger_.warning("revise (merged) rejected due to global structure mismatch, fallback to original")
        return snapshot_text.strip()

    return merged

# 兼容旧入口：默认用分块版
def revise_snapshot_text(snapshot_text: str, comment: str, logger_=None) -> str:
    return revise_snapshot_text_chunked(snapshot_text, comment, logger_, max_chunk_chars=2500)


# ========== DOCX 渲染 ==========
def build_docx_from_snapshot(snapshot_text: str,
                             articles: list[dict],
                             docx_file: str,
                             always_appendix: bool = True,
                             inline_links: bool = True,
                             grouped_for_links: dict | None = None) -> bool:
    """
    兼容两种分区标题：
      1) “综合要闻：”  2) “一、综合要闻：”
    对正文中“单行 URL”渲染为可点击链接；末尾附录列出所有 articles。
    """
    if not snapshot_text:
        return False

    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    lines = snapshot_text.splitlines()
    # 标题
    title = lines[0].strip() if lines else f"中核日报（{datetime.now().strftime('%Y-%m-%d')}）"
    H1 = doc.add_heading(level=1)
    H1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = H1.add_run(title)
    run.font.name = u'宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    url_line_pat = re.compile(r"^(https?://[^\s]+)$")

    i = 1
    while i < len(lines):
        stripped = (lines[i] or "").strip()
        if not stripped:
            i += 1
            continue

        # 分区标题
        m_sec_num = re.match(r"^[一二三四五六七]、(.+?)：$", stripped)
        m_sec_plain = re.match(r"^(.+?)：$", stripped)
        if m_sec_num or m_sec_plain:
            sec_title_cn = (m_sec_num.group(1) if m_sec_num else m_sec_plain.group(1)).strip()
            h = doc.add_heading(level=2)
            rr = h.add_run(f"{sec_title_cn}：")
            rr.font.name = u'宋体'
            rr._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            i += 1
            continue

        # 子类
        if re.match(r"^（(.+?)）$", stripped):
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.bold = True
            i += 1
            continue

        # 编号标题行
        if re.match(r"^\d+，", stripped):
            p = doc.add_paragraph()
            p.add_run(stripped)
            i += 1
            continue

        # 单行 URL → 超链接
        if url_line_pat.match(stripped):
            p2 = doc.add_paragraph()
            add_hyperlink(p2, stripped, stripped)
            i += 1
            continue

        # 普通段落
        doc.add_paragraph(stripped)
        i += 1

    # 文末附录
    if always_appendix:
        doc.add_heading("附：原始信息网页", level=2)
        for k, a in enumerate(articles or [], start=1):
            title_a = a.get("title", "")
            url_a = a.get("url", "")
            d = _norm_date(a.get("publish_time", ""))
            doc.add_paragraph(f"{k}、{title_a}|{d}")
            p2 = doc.add_paragraph()
            if url_a:
                add_hyperlink(p2, url_a, url_a)

    doc.save(docx_file)
    return True
